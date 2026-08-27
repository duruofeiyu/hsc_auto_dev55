from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 打开文档
doc = Document("/Users/a1-6/测试工程师进阶路线图_AI测试开发.docx")

# 添加新章节：2026年8月进展更新
doc.add_page_break()

# 大标题
heading = doc.add_heading("2026年8月进展更新", level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 一、本月完成情况
doc.add_heading("一、本月完成情况", level=2)

doc.add_paragraph("原8月目标：数据驱动（YAML）+ Git/GitHub上传", style="List Bullet")
doc.add_paragraph("实际完成情况：", style="List Bullet")

completed_items = [
    "YAML数据驱动：已完成",
    "ui-testcase-generator Skill调优与使用：已完成",
    "平台设置页33条手工测试用例生成：已完成",
    "发现3个功能缺陷并整理缺陷报告：已完成",
    "框架源码学习（config.py、base.py详解）：已完成",
    "Git/GitHub上传：进行中，预计8月底完成",
]

for item in completed_items:
    doc.add_paragraph(item, style="List Bullet 2")

# 二、新增项目成果
doc.add_heading("二、新增项目成果", level=2)

doc.add_heading("项目1：HSC平台设置页回归测试", level=3)
doc.add_paragraph("时间：2026年8月", style="List Bullet")
doc.add_paragraph("环境：HSC 55开发环境", style="List Bullet")
doc.add_paragraph("内容：", style="List Bullet")
doc.add_paragraph("生成33条手工测试用例（Excel）", style="List Bullet 2")
doc.add_paragraph("覆盖功能测试、兼容测试、性能测试、安全测试", style="List Bullet 2")
doc.add_paragraph("发现1个服务器同步功能缺陷", style="List Bullet 2")
doc.add_paragraph("产出：reports/平台设置页测试用例.xlsx", style="List Bullet")

doc.add_heading("项目2：HSC角色管理缺陷发现", level=3)
doc.add_paragraph("时间：2026年8月", style="List Bullet")
doc.add_paragraph("环境：HSC 55开发环境", style="List Bullet")
doc.add_paragraph("内容：", style="List Bullet")
doc.add_paragraph("发现分配用户/取消授权时空状态提示异常", style="List Bullet 2")
doc.add_paragraph("整理2个缺陷报告", style="List Bullet 2")
doc.add_paragraph("产出：04_踩坑记录/功能缺陷记录.md", style="List Bullet")

# 三、可写进简历的项目更新
doc.add_heading("三、可写进简历的项目更新", level=2)

resume_items = [
    "HSC系统接口自动化测试框架：pytest + requests + Allure + YAML数据驱动，覆盖系统管理5个模块",
    "HSC平台设置页回归测试：生成33条手工测试用例，发现1个前端状态不一致缺陷",
    "HSC角色管理缺陷发现：发现2个空状态操作提示异常缺陷",
    "CI/CD流水线：待完成：GitHub Actions自动跑测试",
]

for item in resume_items:
    doc.add_paragraph(item, style="List Bullet")

# 四、风险与备选方案更新
doc.add_heading("四、风险与备选方案更新", level=2)

doc.add_paragraph("新增观察：", style="List Bullet")
doc.add_paragraph("近期测试中频繁遇到前端提示与实际状态不一致类bug：", style="List Bullet 2")
doc.add_paragraph("平台设置-服务器同步：添加/删除提示与列表状态不一致", style="List Bullet 3")
doc.add_paragraph("角色管理-分配用户：未选用户提示授权成功", style="List Bullet 3")
doc.add_paragraph("角色管理-取消授权：无用户时提示取消授权成功", style="List Bullet 3")

doc.add_paragraph("影响：这类问题可能反映前端状态管理或接口响应处理存在共性问题。", style="List Bullet 2")
doc.add_paragraph("备选：后续测试类似功能时，务必通过刷新页面 + 接口响应双重验证，不轻信前端提示。", style="List Bullet 2")

# 五、下周计划调整
doc.add_heading("五、下周计划调整", level=2)

# 创建表格
table = doc.add_table(rows=3, cols=3)
table.style = "Light Grid Accent 1"

# 表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "时间"
hdr_cells[1].text = "原计划"
hdr_cells[2].text = "调整后"

# 第一行数据
row1_cells = table.rows[1].cells
row1_cells[0].text = "2026.08 第四周"
row1_cells[1].text = "Git/GitHub上传"
row1_cells[2].text = "Git/GitHub上传 + 继续学习utils_*.py接口封装层"

# 第二行数据
row2_cells = table.rows[2].cells
row2_cells[0].text = "2026.09 第一周"
row2_cells[1].text = "GitHub Actions CI/CD"
row2_cells[2].text = "不变"

# 保存
doc.save("/Users/a1-6/测试工程师进阶路线图_AI测试开发.docx")
print("docx 已更新，新增'2026年8月进展更新'章节（不含学习笔记部分）")
