from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import copy

# 打开文档
doc = Document("/Users/a1-6/测试工程师进阶路线图_AI测试开发_备份.docx")

# ==================== 1. 更新月度里程碑表格（Table 1）====================
# 表格索引 0 是第一个表格：时间 | 学习内容 | 产出（写进简历）
table1 = doc.tables[0]
for row in table1.rows:
    cells = row.cells
    if cells[0].text.strip().startswith("2026.08"):
        cells[1].text = "数据驱动（YAML）+ ui-testcase-generator Skill 调优 + 框架源码学习"
        cells[2].text = "代码上传 GitHub（进行中），33 条手工用例，3 个缺陷报告，数据与代码分离"
        break

# ==================== 2. 填充 5.6 月度里程碑（新增）小节 ====================
# 找到 "5.6 月度里程碑（新增）" 这个标题，在它后面插入内容
insert_after_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "5.6 月度里程碑（新增）":
        insert_after_idx = i
        break

if insert_after_idx is not None:
    # 在段落后添加表格
    milestone_table = doc.add_table(rows=5, cols=3)
    milestone_table.style = "Light Grid Accent 1"

    # 表头
    hdr = milestone_table.rows[0].cells
    hdr[0].text = "月份"
    hdr[1].text = "里程碑目标"
    hdr[2].text = "验收标准"

    # 7月
    r1 = milestone_table.rows[1].cells
    r1[0].text = "2026.07"
    r1[1].text = "pytest + requests + fixture + Allure 进阶"
    r1[2].text = "HSC 接口自动化框架跑通，676 条用例 ✅"

    # 8月
    r2 = milestone_table.rows[2].cells
    r2[0].text = "2026.08"
    r2[1].text = "YAML 数据驱动 + Skill 调优 + 框架源码理解 + GitHub 上传"
    r2[2].text = "33 条手工用例 + 3 个缺陷报告 + config.py/base.py 源码笔记 + GitHub 仓库"

    # 9月
    r3 = milestone_table.rows[3].cells
    r3[0].text = "2026.09"
    r3[1].text = "GitHub Actions CI/CD 入门"
    r3[2].text = "提交代码自动跑测试，生成 Allure 报告"

    # 10-12月
    r4 = milestone_table.rows[4].cells
    r4[0].text = "2026.10-12"
    r4[1].text = "Playwright UI 自动化（选修）+ 简历准备 + 秋招"
    r4[2].text = "完整 GitHub 仓库 + 秋招简历 + 面试准备"

    # 把新添加的表格移动到指定位置后面
    # docx 中表格和段落是混合的，这里用简单方法：插入一个标记段落再替换
    # 更简单的方法：直接在标题后面添加段落说明
    p = doc.paragraphs[insert_after_idx]
    p._element.addnext(milestone_table._element)

# ==================== 3. 更新"当前最优先的任务" ====================
# 找到对应的段落并替换
for para in doc.paragraphs:
    if para.text.strip().startswith("本周：整理 HSC 代码结构"):
        para.text = "本周：完成平台设置页回归测试，整理框架源码笔记，调优 ui-testcase-generator skill"
    elif para.text.strip().startswith("下周：把整理好的代码上传到 GitHub"):
        para.text = "下周：把整理好的代码上传到 GitHub，写好 README；继续学习 utils_*.py 接口封装层"
    elif para.text.strip().startswith("下下周：接入 GitHub Actions"):
        para.text = "下下周：接入 GitHub Actions，实现提交代码自动跑测试"
    elif para.text.strip().startswith("最后更新：2026 年 7 月 31 日"):
        para.text = "最后更新：2026 年 8 月 10 日"

# ==================== 4. 更新"可写进简历的项目建议"数字 ====================
for para in doc.paragraphs:
    if "基于 PRD 输出 60+ 条手工测试用例" in para.text:
        para.text = para.text.replace("基于 PRD 输出 60+ 条手工测试用例", "基于 PRD/需求输出 60+ 条手工测试用例（平台设置页 33 条）")
    if "发现并跟踪 5+ 后端缺陷" in para.text:
        para.text = para.text.replace("发现并跟踪 5+ 后端缺陷", "发现并跟踪 5+ 功能缺陷（含前端状态不一致类问题）")

# ==================== 5. 添加 8 月进展更新章节（和之前一样）====================
doc.add_page_break()

heading = doc.add_heading("2026年8月进展更新", level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

doc.add_heading("一、本月完成情况", level=2)
doc.add_paragraph("原8月目标：数据驱动（YAML）+ Git/GitHub上传", style="List Bullet")
doc.add_paragraph("实际完成情况：", style="List Bullet")
completed_items = [
    "YAML数据驱动：已完成",
    "ui-testcase-generator Skill调优与使用：已完成",
    "平台设置页33条手工测试用例生成：已完成",
    "发现3个功能缺陷并整理缺陷报告：已完成",
    "框架源码学习（config.py、base.py详解）：已完成",
    "Git/GitHub上传：进行中，预计8月底完成",
]
for item in completed_items:
    doc.add_paragraph(item, style="List Bullet 2")

doc.add_heading("二、新增项目成果", level=2)
doc.add_heading("项目1：HSC平台设置页回归测试", level=3)
doc.add_paragraph("时间：2026年8月", style="List Bullet")
doc.add_paragraph("环境：HSC 55开发环境", style="List Bullet")
doc.add_paragraph("内容：", style="List Bullet")
doc.add_paragraph("生成33条手工测试用例（Excel）", style="List Bullet 2")
doc.add_paragraph("覆盖功能测试、兼容测试、性能测试、安全测试", style="List Bullet 2")
doc.add_paragraph("发现1个服务器同步功能缺陷", style="List Bullet 2")
doc.add_paragraph("产出：reports/平台设置页测试用例.xlsx", style="List Bullet")

doc.add_heading("项目2：HSC角色管理缺陷发现", level=3)
doc.add_paragraph("时间：2026年8月", style="List Bullet")
doc.add_paragraph("环境：HSC 55开发环境", style="List Bullet")
doc.add_paragraph("内容：", style="List Bullet")
doc.add_paragraph("发现分配用户/取消授权时空状态提示异常", style="List Bullet 2")
doc.add_paragraph("整理2个缺陷报告", style="List Bullet 2")
doc.add_paragraph("产出：04_踩坑记录/功能缺陷记录.md", style="List Bullet")

doc.add_heading("三、可写进简历的项目更新", level=2)
resume_items = [
    "HSC系统接口自动化测试框架：pytest + requests + Allure + YAML数据驱动，覆盖系统管理5个模块",
    "HSC平台设置页回归测试：生成33条手工测试用例，发现1个前端状态不一致缺陷",
    "HSC角色管理缺陷发现：发现2个空状态操作提示异常缺陷",
    "CI/CD流水线：待完成：GitHub Actions自动跑测试",
]
for item in resume_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("四、风险与备选方案更新", level=2)
doc.add_paragraph("新增观察：", style="List Bullet")
doc.add_paragraph("近期测试中频繁遇到前端提示与实际状态不一致类bug：", style="List Bullet 2")
doc.add_paragraph("平台设置-服务器同步：添加/删除提示与列表状态不一致", style="List Bullet 3")
doc.add_paragraph("角色管理-分配用户：未选用户提示授权成功", style="List Bullet 3")
doc.add_paragraph("角色管理-取消授权：无用户时提示取消授权成功", style="List Bullet 3")
doc.add_paragraph("影响：这类问题可能反映前端状态管理或接口响应处理存在共性问题。", style="List Bullet 2")
doc.add_paragraph("备选：后续测试类似功能时，务必通过刷新页面 + 接口响应双重验证，不轻信前端提示。", style="List Bullet 2")

doc.add_heading("五、下周计划调整", level=2)
plan_table = doc.add_table(rows=3, cols=3)
plan_table.style = "Light Grid Accent 1"
hdr = plan_table.rows[0].cells
hdr[0].text = "时间"
hdr[1].text = "原计划"
hdr[2].text = "调整后"
r1 = plan_table.rows[1].cells
r1[0].text = "2026.08 第四周"
r1[1].text = "Git/GitHub上传"
r1[2].text = "Git/GitHub上传 + 继续学习utils_*.py接口封装层"
r2 = plan_table.rows[2].cells
r2[0].text = "2026.09 第一周"
r2[1].text = "GitHub Actions CI/CD"
r2[2].text = "不变"

# 保存
doc.save("/Users/a1-6/测试工程师进阶路线图_AI测试开发.docx")
print("docx 已全面更新")
print("主要更新：")
print("1. 月度里程碑表格（Table 1）2026.08 内容更新")
print("2. 5.6 月度里程碑小节填充表格")
print("3. 当前最优先任务更新为当前实际")
print("4. 最后更新日期改为 2026 年 8 月 10 日")
print("5. 可写进简历的项目建议描述细化")
print("6. 末尾新增 2026年8月进展更新章节")
