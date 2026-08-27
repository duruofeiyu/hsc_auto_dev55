from docx import Document
from docx.oxml import OxmlElement

# 打开已更新的文档
doc = Document("/Users/a1-6/测试工程师进阶路线图_AI测试开发.docx")

# ==================== 1. 删除错误放置的月度里程碑表格 ====================
# 找到表头为 "月份 | 里程碑目标 | 验收标准" 的表格并删除
table_to_remove = None
for table in doc.tables:
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    if header_cells == ["月份", "里程碑目标", "验收标准"]:
        table_to_remove = table
        break

if table_to_remove:
    # 删除表格的 XML 元素
    table_to_remove._element.getparent().remove(table_to_remove._element)
    print("已删除错误放置的月度里程碑表格")

# ==================== 2. 在 5.6 标题后插入里程碑内容（用段落形式）====================
insert_after_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "5.6 月度里程碑（新增）":
        insert_after_idx = i
        break

if insert_after_idx is not None:
    # 在 5.6 标题后面插入内容
    target_para = doc.paragraphs[insert_after_idx]

    # 使用 _element.addnext 在段落后添加新段落
    new_items = [
        ("2026.07：pytest + requests + fixture + Allure 进阶，完成 HSC 接口自动化框架（676 条用例）。", "List Bullet"),
        ("2026.08：YAML 数据驱动 + ui-testcase-generator Skill 调优 + 框架源码理解 + GitHub 仓库上传。验收：33 条手工用例 + 3 个缺陷报告 + config.py/base.py 源码笔记。", "List Bullet"),
        ("2026.09：GitHub Actions CI/CD 入门，实现提交代码自动跑测试并生成 Allure 报告。", "List Bullet"),
        ("2026.10-12：Playwright UI 自动化（选修）+ 简历准备 + 秋招投递。", "List Bullet"),
    ]

    # 需要倒序插入，因为每次 addnext 都是在 target 后面插入，后插入的会排在后面
    for text, style in reversed(new_items):
        new_p = doc.add_paragraph(text, style=style)
        target_para._element.addnext(new_p._element)

    print("已在 5.6 标题后插入月度里程碑内容")

# 保存
doc.save("/Users/a1-6/测试工程师进阶路线图_AI测试开发.docx")
print("docx 修复完成")
